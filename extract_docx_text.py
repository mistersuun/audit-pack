import docx
import sys

def extract_text_from_docx(docx_path, output_txt_path):
    try:
        doc = docx.Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_text))

        with open(output_txt_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))
            
        print(f"Successfully extracted text to {output_txt_path}")
    except Exception as e:
        print(f"Error extracting text: {e}")
        sys.exit(1)

if __name__ == "__main__":
    input_file = "documentation/back/2025-02 - Procédure Complete Back (Audition).docx"
    output_file = "documentation/back/procedure_extracted.txt"
    extract_text_from_docx(input_file, output_file)
