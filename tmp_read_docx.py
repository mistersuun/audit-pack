import sys
import os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

base = r'C:\Users\Auditeur\Documents\Projects\audit-pack'
files = [
    'RJ_Audit_Methodology.docx',
    'RJ_Balancing_Complete_Guide.docx',
    'RJ_Balancing_Guide_Sheraton.docx',
]

for f in files:
    path = os.path.join(base, f)
    out_path = os.path.join(base, f.replace('.docx', '_EXTRACT.txt'))
    print(f'Extracting {f} -> {out_path}')
    try:
        doc = Document(path)
        with open(out_path, 'w', encoding='utf-8') as fh:
            fh.write(f'=== {f} ===\n\n')
            for p in doc.paragraphs:
                if p.text.strip():
                    fh.write(p.text + '\n')
            for ti, table in enumerate(doc.tables):
                fh.write(f'\n[TABLE {ti}]\n')
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    fh.write(' | '.join(cells) + '\n')
        print(f'  size: {os.path.getsize(out_path)}')
    except Exception as e:
        print(f'  ERROR: {e}')
