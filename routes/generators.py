"""
Blueprint for generating night audit template documents dynamically.

Each generator loads the ACTUAL template file and replaces only the dynamic
fields (dates, etc.) while preserving all original formatting, layout,
fonts, margins, and page orientation.
"""

from flask import Blueprint, request, jsonify, send_file, render_template
from functools import wraps
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from openpyxl import load_workbook
from lxml import etree
import copy
import re
import io
import os
from utils.auth_decorators import login_required
from utils.weather_capture import get_weather_screenshot, fetch_tomorrow_weather
import logging

logger = logging.getLogger(__name__)

generators_bp = Blueprint('generators', __name__)

@generators_bp.route('/generators')
@login_required
def generators_page():
    """Display the generators page."""
    return render_template('generators.html')

# Path to template files (absolute, relative to project root)
TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'static', 'templates')

# French day/month names (shared across generators)
DAYS_FR = ['Lundi', 'Mardi', 'Mercredi', 'Jeudi', 'Vendredi', 'Samedi', 'Dimanche']
MONTHS_FR = ['Janvier', 'Février', 'Mars', 'Avril', 'Mai', 'Juin',
             'Juillet', 'Août', 'Septembre', 'Octobre', 'Novembre', 'Décembre']

# Word XML namespace
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Known banquet / meeting rooms at Sheraton Laval
SALONS = [
    'Auteuil',
    'Vimont',
    'Duvernay',
    'Chomedey',
    'Ste-Dorothée',
    'Lavaloise',
    'Laval 1',
    'Laval 2',
    'Laval 3',
    'Laval 4',
    'Laval 5',
    'Giotto 1',
    'Giotto 2',
    'Giotto 3',
    'Giotto 4',
    'Giotto 5',
    'Centre des congrès',
    'Foyer Laval',
]


def _format_date_fr(date_obj):
    """Format a date in French: 'Lundi le 17 Novembre 2025'"""
    day_name = DAYS_FR[date_obj.weekday()]
    month_name = MONTHS_FR[date_obj.month - 1]
    return f"{day_name} le {date_obj.day} {month_name} {date_obj.year}"


def _docx_to_pdf(docx_bytes):
    """Convert .docx bytes to PDF bytes using Microsoft Word via COM.

    Requires Microsoft Word installed on the machine.
    Returns PDF bytes or None on failure.
    """
    import tempfile
    try:
        import win32com.client
        import pythoncom
    except ImportError:
        return None

    docx_path = None
    pdf_path = None
    try:
        pythoncom.CoInitialize()

        # Write .docx to temp file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.write(docx_bytes)
            docx_path = tmp.name

        pdf_path = docx_path.replace('.docx', '.pdf')

        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        word.DisplayAlerts = False

        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
        doc.Close(False)
        word.Quit()

        with open(pdf_path, 'rb') as f:
            return f.read()
    except Exception as e:
        logger.error(f"DOCX to PDF conversion failed: {e}")
        return None
    finally:
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass
        for p in [docx_path, pdf_path]:
            if p and os.path.exists(p):
                try:
                    os.unlink(p)
                except Exception:
                    pass


def _xlsx_to_pdf(xlsx_bytes):
    """Convert .xlsx bytes to PDF bytes using Microsoft Excel via COM."""
    import tempfile
    try:
        import win32com.client
        import pythoncom
    except ImportError:
        return None

    xlsx_path = None
    pdf_path = None
    try:
        pythoncom.CoInitialize()

        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            tmp.write(xlsx_bytes)
            xlsx_path = tmp.name

        pdf_path = xlsx_path.replace('.xlsx', '.pdf')

        excel = win32com.client.Dispatch('Excel.Application')
        excel.Visible = False
        excel.DisplayAlerts = False

        wb = excel.Workbooks.Open(os.path.abspath(xlsx_path))
        wb.ExportAsFixedFormat(0, os.path.abspath(pdf_path))  # 0 = xlTypePDF
        wb.Close(False)
        excel.Quit()

        with open(pdf_path, 'rb') as f:
            return f.read()
    except Exception as e:
        logger.error(f"XLSX to PDF conversion failed: {e}")
        return None
    finally:
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass
        for p in [xlsx_path, pdf_path]:
            if p and os.path.exists(p):
                try:
                    os.unlink(p)
                except Exception:
                    pass


def _replace_runs_text(paragraph, new_text, leading_spaces=0):
    """
    Replace all text in a paragraph's runs with new_text while preserving
    the formatting (font size, bold, color, lang) of the first run.
    Clears all existing runs and creates a single new one.
    """
    runs = paragraph.runs
    if not runs:
        return

    first_run_elem = runs[0]._element
    rpr = first_run_elem.find('{%s}rPr' % W_NS)
    rpr_copy = copy.deepcopy(rpr) if rpr is not None else None

    p_elem = paragraph._element
    for r in p_elem.findall('{%s}r' % W_NS):
        p_elem.remove(r)

    new_run = etree.SubElement(p_elem, '{%s}r' % W_NS)
    if rpr_copy is not None:
        new_run.insert(0, rpr_copy)
    new_t = etree.SubElement(new_run, '{%s}t' % W_NS)
    new_t.text = (' ' * leading_spaces) + new_text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


# ─────────────────────────────────────────────────────────────────────────────
# PRINT PREVIEW — renders printable HTML for any generator
# ─────────────────────────────────────────────────────────────────────────────

@generators_bp.route('/api/generators/print/<gen_type>', methods=['GET'])
@login_required
def print_preview(gen_type):
    """
    Generate the same document as 'Generer', convert to PDF via Word/Excel COM,
    and serve it to the browser for instant print preview.
    Query param: ?date=YYYY-MM-DD
    """
    date_str = request.args.get('date')
    if not date_str:
        return jsonify({'error': 'Date manquante'}), 400

    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
    except ValueError:
        return jsonify({'error': 'Format de date invalide'}), 400

    # Map gen_type to the generator function and file type
    gen_map = {
        'separateur-date': ('_gen_separateur', 'docx'),
        'checklist-tournee': ('_gen_checklist_tournee', 'xlsx'),
        'entretien-hiver': ('_gen_entretien', 'docx'),
    }

    if gen_type not in gen_map:
        return jsonify({'error': 'Type inconnu'}), 404

    func_name, file_type = gen_map[gen_type]

    try:
        # Generate the document bytes using the same logic as 'Generer'
        doc_bytes = globals()[func_name](date_obj)

        # Convert to PDF
        if file_type == 'docx':
            pdf_bytes = _docx_to_pdf(doc_bytes)
        else:
            pdf_bytes = _xlsx_to_pdf(doc_bytes)

        if pdf_bytes:
            return send_file(
                io.BytesIO(pdf_bytes),
                mimetype='application/pdf',
                as_attachment=False,
                download_name=f'{gen_type}_{date_str}.pdf'
            )
        else:
            # Fallback: serve the original file for download
            mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' if file_type == 'docx' else 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            return send_file(
                io.BytesIO(doc_bytes),
                mimetype=mime,
                as_attachment=True,
                download_name=f'{gen_type}_{date_str}.{file_type}'
            )
    except Exception as e:
        logger.exception(f"Print preview error for {gen_type}")
        return jsonify({'error': str(e)}), 500


def _gen_separateur(date_obj):
    """Generate separateur date .docx bytes."""
    new_date_text = _format_date_fr(date_obj)
    template_path = os.path.join(TEMPLATE_DIR, 'separateur_date_comptabilité.docx')
    doc = Document(template_path)
    body = doc.element.body
    txbx_elements = body.findall('.//' + '{%s}txbxContent' % W_NS)
    for txbx in txbx_elements:
        paragraphs = txbx.findall('{%s}p' % W_NS)
        if not paragraphs:
            continue
        p_elem = paragraphs[0]
        runs = p_elem.findall('{%s}r' % W_NS)
        if not runs:
            continue
        # Replace text in first run, clear others
        for i, r in enumerate(runs):
            t = r.find('{%s}t' % W_NS)
            if t is not None:
                if i == 0:
                    t.text = new_date_text
                else:
                    t.text = ''
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _gen_checklist_tournee(date_obj):
    """Generate checklist tournee .xlsx bytes."""
    template_path = os.path.join(TEMPLATE_DIR, 'Checklist Tournée Étages.xlsx')
    wb = load_workbook(template_path)
    ws = wb.active
    ws['B2'] = date_obj.strftime('%m/%d/%Y')
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _gen_entretien(date_obj):
    """Generate entretien hiver .docx bytes."""
    day_name = DAYS_FR[date_obj.weekday()]
    day_num = date_obj.day
    month_name = MONTHS_FR[date_obj.month - 1]
    year = date_obj.year
    if date_obj.month >= 10:
        season = f"{year}-{year+1}"
    else:
        season = f"{year-1}-{year}"

    template_path = os.path.join(TEMPLATE_DIR, 'Entretien Sheraton Laval Hiver.docx')
    doc = Document(template_path)

    replacements = {
        '2024-2025': season,
        'Vendredi': day_name,
        '17': str(day_num),
        'Novembre': month_name,
        '2025': str(year),
    }
    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                for run in p.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# API: Salon list for dropdown
# ─────────────────────────────────────────────────────────────────────────────

@generators_bp.route('/api/generators/salons', methods=['GET'])
@login_required
def get_salons():
    """Return the list of known banquet/meeting room names."""
    return jsonify(SALONS)


@generators_bp.route('/api/generators/weather', methods=['GET'])
@login_required
def get_weather_preview():
    """Fetch tomorrow's weather for Laval via Open-Meteo (no API key needed)."""
    data = fetch_tomorrow_weather()
    if data is None:
        return jsonify({'error': 'Météo indisponible'}), 503
    # Remove non-serializable date_obj
    result = {k: v for k, v in data.items() if k != 'date_obj'}
    return jsonify(result)


# ─────────────────────────────────────────────────────────────────────────────
# 1. SÉPARATEUR DATE
# ─────────────────────────────────────────────────────────────────────────────

@generators_bp.route('/api/generators/separateur-date', methods=['POST'])
@login_required
def generate_separateur_date():
    """
    Generate Séparateur Date document from actual template.
    The template has a VML text box with "Lundi le 17 Novembre 2025" at 62pt
    centered in landscape orientation. We replace the text in the textbox
    while preserving all formatting.
    """
    data = request.get_json()
    date_str = data.get('date')

    if not date_str:
        return jsonify({'error': 'Date is required'}), 400

    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        new_date_text = _format_date_fr(date_obj)

        template_path = os.path.join(TEMPLATE_DIR, 'separateur_date_comptabilité.docx')
        doc = Document(template_path)

        body = doc.element.body

        # Keep original VML text box position (right half, at page edge).
        # Template: margin-left:367.5pt; width:378pt — no changes needed.
        VML_NS = 'urn:schemas-microsoft-com:vml'

        # Replace text in VML textbox
        txbx_elements = body.findall('.//' + '{%s}txbxContent' % W_NS)

        for txbx in txbx_elements:
            paragraphs = txbx.findall('{%s}p' % W_NS)
            if not paragraphs:
                continue

            p_elem = paragraphs[0]
            runs = p_elem.findall('{%s}r' % W_NS)
            if not runs:
                continue

            existing_text = ''
            for r in runs:
                t_elem = r.find('{%s}t' % W_NS)
                if t_elem is not None and t_elem.text:
                    existing_text += t_elem.text

            if 'le' not in existing_text.lower():
                continue

            first_run = runs[0]
            rpr = first_run.find('{%s}rPr' % W_NS)
            rpr_copy = copy.deepcopy(rpr) if rpr is not None else None

            for r in runs:
                p_elem.remove(r)

            new_run = etree.SubElement(p_elem, '{%s}r' % W_NS)
            if rpr_copy is not None:
                new_run.insert(0, rpr_copy)
            new_t = etree.SubElement(new_run, '{%s}t' % W_NS)
            new_t.text = new_date_text
            new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f'Separateur_Date_{date_obj.strftime("%Y-%m-%d")}.docx'
        return send_file(buffer, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except ValueError:
        return jsonify({'error': 'Format de date invalide. Utiliser AAAA-MM-JJ'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ─────────────────────────────────────────────────────────────────────────────
# 2. CHECKLIST TOURNÉE DES ÉTAGES
# ─────────────────────────────────────────────────────────────────────────────

@generators_bp.route('/api/generators/checklist-tournee', methods=['POST'])
@login_required
def generate_checklist_tournee():
    """
    Generate Checklist Tournée des Étages from actual Excel template.
    Date in merged cell A3 (A3:F3), Arial 22pt, centered.
    """
    data = request.get_json()
    date_str = data.get('date')

    if not date_str:
        return jsonify({'error': 'Date is required'}), 400

    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        formatted_date = date_obj.strftime('%m/%d/%Y')

        template_path = os.path.join(TEMPLATE_DIR, 'Checklist Tournée Étages.xlsx')
        wb = load_workbook(template_path)
        ws = wb.active

        ws['A3'] = f"liste des vérifications pour auditeur de nuit {formatted_date}"

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        filename = f'Checklist_Tournee_{date_obj.strftime("%Y-%m-%d")}.xlsx'
        return send_file(buffer, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

    except ValueError:
        return jsonify({'error': 'Format de date invalide. Utiliser AAAA-MM-JJ'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ─────────────────────────────────────────────────────────────────────────────
# 3. ENTRETIEN SHERATON LAVAL HIVER
# ─────────────────────────────────────────────────────────────────────────────

@generators_bp.route('/api/generators/entretien-hiver', methods=['POST'])
@login_required
def generate_entretien_hiver():
    """
    Generate Feuille d'Entretien from actual Word template.
    Updates season year + single date in table header.
    Captures weather screenshot for tomorrow only.
    """
    data = request.get_json()
    date_str = data.get('date')

    if not date_str:
        return jsonify({'error': 'Date is required'}), 400

    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')

        day_name = DAYS_FR[date_obj.weekday()]
        day_num = date_obj.day
        month_name = MONTHS_FR[date_obj.month - 1]
        year = date_obj.year

        # Determine winter season
        if date_obj.month >= 10:
            season_start, season_end = year, year + 1
        else:
            season_start, season_end = year - 1, year

        template_path = os.path.join(TEMPLATE_DIR, 'Entretien Sheraton Laval Hiver.docx')
        doc = Document(template_path)

        if doc.tables:
            table = doc.tables[0]
            cell = table.rows[0].cells[0]

            if len(cell.paragraphs) >= 1:
                _replace_runs_text(cell.paragraphs[0],
                                   f"Entretien Sheraton Laval hiver {season_start}-{season_end}",
                                   leading_spaces=4)

            if len(cell.paragraphs) >= 2:
                new_date_line = f"{day_name} {day_num} {month_name} {year}"
                _replace_runs_text(cell.paragraphs[1], new_date_line, leading_spaces=52)

        # Remove old weather screenshot images from the template.
        # They live in body paragraphs that contain <w:drawing> elements.
        body = doc.element.body
        W_NS_URI = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        paras_to_remove = []
        for child in body:
            if child.tag == f'{{{W_NS_URI}}}p':
                if child.findall('.//{%s}drawing' % W_NS_URI):
                    paras_to_remove.append(child)
        for p_elem in paras_to_remove:
            body.remove(p_elem)

        # Generate and insert new weather card for the selected date
        print(f"Capturing weather forecast for {date_str}...")
        weather_image = get_weather_screenshot(date_str)

        if weather_image:
            img_buffer = io.BytesIO()
            weather_image.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            wp = doc.add_paragraph()
            wp.add_run().add_picture(img_buffer, width=Inches(10.0))
        else:
            wp = doc.add_paragraph()
            wp.add_run("Prévisions météo temporairement indisponibles").bold = True

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f'Entretien_Hiver_{date_obj.strftime("%Y-%m-%d")}.docx'
        return send_file(buffer, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except ValueError:
        return jsonify({'error': 'Format de date invalide. Utiliser AAAA-MM-JJ'}), 400
    except Exception as e:
        print(f"Error generating document: {e}")
        return jsonify({'error': str(e)}), 500


# ─────────────────────────────────────────────────────────────────────────────
# 4. CLÉS BANQUETS
# ─────────────────────────────────────────────────────────────────────────────

@generators_bp.route('/api/generators/cles-banquets', methods=['POST'])
@login_required
def generate_cles_banquets():
    """
    Generate Clés Banquets document with event details.
    Built from scratch (original .doc template is corrupted).
    """
    data = request.get_json()
    date_str = data.get('date')
    events = data.get('events', [])

    if not date_str:
        return jsonify({'error': 'Date is required'}), 400

    if not events:
        return jsonify({'error': 'Au moins un événement est requis'}), 400

    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        formatted_date = _format_date_fr(date_obj)

        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        title_p = doc.add_paragraph()
        title_p.alignment = 1
        run = title_p.add_run('CLÉS BANQUETS')
        run.font.size = Pt(20)
        run.font.bold = True

        date_p = doc.add_paragraph()
        date_p.alignment = 1
        run = date_p.add_run(f'DATE : {formatted_date}')
        run.font.size = Pt(14)
        run.font.bold = True

        doc.add_paragraph()

        headers = ['Salon', 'Compagnie', 'Heure', 'Responsable', 'Signature']
        table = doc.add_table(rows=1 + len(events), cols=5)
        table.style = 'Table Grid'

        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for r in cell.paragraphs[0].runs:
                r.font.bold = True
                r.font.size = Pt(11)

        for idx, event in enumerate(events):
            row = table.rows[idx + 1]
            row.cells[0].text = event.get('salon', '')
            row.cells[1].text = event.get('compagnie', '')
            row.cells[2].text = event.get('heure', '')
            row.cells[3].text = event.get('responsable', '')

            for cell in row.cells:
                for r in cell.paragraphs[0].runs:
                    r.font.size = Pt(11)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f'Cles_Banquets_{date_obj.strftime("%Y-%m-%d")}.docx'
        return send_file(buffer, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except ValueError:
        return jsonify({'error': 'Format de date invalide. Utiliser AAAA-MM-JJ'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ─────────────────────────────────────────────────────────────────────────────
# 5. CARGO JET — Crew Room Record
# ─────────────────────────────────────────────────────────────────────────────

@generators_bp.route('/api/generators/cargo-jet', methods=['POST'])
@login_required
def generate_cargo_jet():
    """
    Generate Cargo Jet crew room record from Excel template.
    Row 6: logo + date | Row 7-8: headers
    Row 9+: crew data (DATE, TIME, PRINT NAME, BonVoy, SIGNATURE, ROOM, DATE-out, TIME-out)
    """
    data = request.get_json()
    date_str = data.get('date')
    crew = data.get('crew', [])

    if not date_str:
        return jsonify({'error': 'Date is required'}), 400

    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')

        template_path = os.path.join(TEMPLATE_DIR, 'CARGO JET.xlsx')
        wb = load_workbook(template_path)
        ws = wb['Feuil1']

        # Update the date in H6
        ws['H6'] = date_obj

        # Clear existing data rows (9–22)
        for r in range(9, 23):
            for c in range(2, 10):  # B–I
                ws.cell(r, c).value = None

        # Fill crew
        for i, member in enumerate(crew):
            row = 9 + i
            if row > 22:
                break
            ws.cell(row, 2).value = date_obj                   # B: Check-in date
            ws.cell(row, 4).value = member.get('name', '')      # D: Print name
            ws.cell(row, 7).value = member.get('room', '')      # G: Room

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        filename = f'Cargo_Jet_{date_obj.strftime("%Y-%m-%d")}.xlsx'
        return send_file(buffer, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

    except ValueError:
        return jsonify({'error': 'Format de date invalide. Utiliser AAAA-MM-JJ'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ─────────────────────────────────────────────────────────────────────────────
# 6. FEDEX — Crew Room Record
# ─────────────────────────────────────────────────────────────────────────────

@generators_bp.route('/api/generators/fedex', methods=['POST'])
@login_required
def generate_fedex():
    """
    Generate FedEx crew room record from Excel template.
    Row 6: "FEDEX" + date | Row 7-8: headers
    Row 10+: crew data (DATE, TIME, PRINT NAME, EMPL.NO, FLIGHT, BonVoy, SIGNATURE, ROOM, DATE-out, TIME-out)
    """
    data = request.get_json()
    date_str = data.get('date')
    crew = data.get('crew', [])

    if not date_str:
        return jsonify({'error': 'Date is required'}), 400

    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')

        template_path = os.path.join(TEMPLATE_DIR, 'fedex room record.xlsx')
        wb = load_workbook(template_path)
        ws = wb['Feuil1']

        # Update date
        ws['I6'] = f"Date: {date_obj.strftime('%m/%d/%Y')}"

        # Clear existing data rows (10–18, avoiding merged H19:K20)
        for r in range(10, 19):
            for c in range(2, 12):  # B–K
                ws.cell(r, c).value = None

        # Fill crew (max row 18 to avoid merged H19:K20)
        for i, member in enumerate(crew):
            row = 10 + i
            if row > 18:
                break
            ws.cell(row, 2).value = date_obj                       # B: Check-in date
            ws.cell(row, 4).value = member.get('name', '')          # D: Print name
            ws.cell(row, 5).value = member.get('employee_no', '')   # E: Employee no.
            ws.cell(row, 6).value = member.get('flight', '')        # F: Flight
            ws.cell(row, 9).value = member.get('room', '')          # I: Room

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        filename = f'FedEx_{date_obj.strftime("%Y-%m-%d")}.xlsx'
        return send_file(buffer, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

    except ValueError:
        return jsonify({'error': 'Format de date invalide. Utiliser AAAA-MM-JJ'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500
